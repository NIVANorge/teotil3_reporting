import os

import cartopy.crs as ccrs
import cartopy.io.shapereader as shpreader
import geopandas as gpd
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from matplotlib.transforms import Bbox
from matplotlib_scalebar.scalebar import ScaleBar

plt.style.use("ggplot")

# Bounding box for Norway in ETRS89 UTM Zone 33N
EPSG = 25833
XMIN, XMAX, YMIN, YMAX = -120000, 1180000, 6400000, 8000000

# Colours for report categories
COLOUR_DICT = {
    "Akvakultur": "royalblue",
    "Jordbruk": "sienna",
    "Avløp": "red",
    "Industri": "darkgrey",
    "Bebygd": "gold",
    "Bakgrunn": "limegreen",
    "Menneskeskapt": "pink",
    "Totalt": "black",
}
plt.rcParams["axes.prop_cycle"] = plt.cycler(color=list(COLOUR_DICT.values()))

# Define regions used in report by grouping main catchments. Intervals are "Python-style"
# i.e. include first but not last element in range
REGIONS_DICT = {
    # Definerte kystavsnitt
    "Norges kystområder": [1, 248, 315],
    "Sverige – Strømtangen fyr (Fredrikstad)": [1, 3],
    "Indre Oslofjord (nord for Drøbak)": [5, 10],
    "Hele Oslofjord (Svenskegrensa - Kragerø)": [1, 18],
    "Svenskegrensa – Lindesnes": [1, 24],
    "Lindesnes – Stad": [24, 92],
    "Stad – Russland": [92, 248],
    # Norske vannregioner
    "Glomma": [1, 11],
    "Vest-Viken": [11, 18],
    "Agder": [18, 27],
    "Rogaland": [27, 41],
    "Vestland": [41, 92],
    "Møre og Romsdal": [92, 117],
    "Trøndelag": [117, 144],
    "Nordland": [144, 186],
    "Troms": [186, 211],
    "Finnmark": [211, 248],
    # Norske forvaltingsplanområder
    "Nordsjøen": [1, 91, 315],
    "Norskehavet": [91, 171],
    "Barentshavet": [171, 248],
}


def get_teotil2_results_for_vassoms(st_yr, end_yr, vassom_list):
    """Reads TEOTIL2 results from GitHub.

    Args
        st_yr: Int. First year of interest.
        end_yr: Int. Last year of interest.
        vassom_list: List of Int. List of IDs for vassdragsområder of interest.

    Returns
        Dataframe.
    """
    main_catches = [f"{i:03d}." for i in vassom_list]
    df_list = []
    for year in range(st_yr, end_yr + 1):
        base_url = f"https://raw.githubusercontent.com/NIVANorge/teotil2/main/data/norway_annual_output_data/teotil2_results_{year}.csv"
        df = pd.read_csv(base_url)
        df = df.query("regine in @main_catches").copy()
        df["År"] = year
        cols = [i for i in df.columns if i.split("_")[0] == "accum"]
        df = df[["regine", "År"] + cols]
        df_list.append(df)
    df = pd.concat(df_list)

    return df


def get_teotil3_results_for_vassoms(
    st_yr, end_yr, vassom_list, agri_loss_model, nve_data_yr
):
    """Reads TEOTIL3 results from JupyterHub.

    Args
        st_yr: Int. First year of interest.
        end_yr: Int. Last year of interest.
        vassom_list: List of Int. List of IDs for vassdragsområder of interest.
        agri_loss_model: Str. Either 'risk' or 'annual'.
        nve_data_year: Int. Delivery year of NVE flow dataset to use.

    Returns
        Dataframe.
    """
    if agri_loss_model not in ("risk", "annual"):
        raise ValeError("'agri_loss_model' must be either 'risk' or 'annual'.")

    main_catches = [f"{i:03d}." for i in vassom_list]
    df = pd.read_csv(
        f"/home/jovyan/shared/common/teotil3/evaluation/teo3_results_nve{nve_data_yr}_{st_yr}-{end_yr}_agri-{agri_loss_model}-loss.csv"
    )
    df = df.query(
        "(regine in @main_catches) and (year >= @st_yr) and (year <= @end_yr)"
    ).copy()
    df["År"] = df["year"]
    cols = [i for i in df.columns if i.split("_")[0] == "accum"]
    df = df[["regine", "År"] + cols]
    for col in df.columns:
        if col.endswith("_kg"):
            new_col = col.replace("_kg", "_tonnes")
            df[new_col] = df[col] / 1000
            del df[col]

    return df


def get_aggregation_dict_for_columns(par, model="teotil2"):
    """Make a dict mapping TEOTIL column names to summary columns with
    aggregation where necessary.

    Args
        par: Str. Either 'n' or 'p'.
        model: Str. Either 'teotil2' or 'teotil3'.

    Returns
        Dict with key's equal to output headings and values are lists
        of columns to aggregate.
    """
    valid_models = {"teotil2": ("n", "p"), "teotil3": ("n", "p", "c")}

    if model not in valid_models:
        raise ValueError("'model' must be either 'teotil2' or 'teotil3'.")

    if par not in valid_models[model]:
        raise ValueError(
            f"'par' must be one of {valid_models[model]} for {model.upper()}."
        )

    if model == "teotil2":
        agg_dict = {
            "Akvakultur": [f"accum_aqu_tot-{par}_tonnes"],
            "Jordbruk": [
                f"accum_agri_diff_tot-{par}_tonnes",
                f"accum_agri_pt_tot-{par}_tonnes",
            ],
            "Avløp": [f"accum_ren_tot-{par}_tonnes", f"accum_spr_tot-{par}_tonnes"],
            "Industri": [f"accum_ind_tot-{par}_tonnes"],
            "Bebygd": [f"accum_urban_tot-{par}_tonnes"],
            "Bakgrunn": [f"accum_nat_diff_tot-{par}_tonnes"],
        }
    elif par in ("n", "p"):
        agg_dict = {
            "Akvakultur": [f"accum_aquaculture_tot{par}_tonnes"],
            "Jordbruk": [f"accum_agriculture_tot{par}_tonnes"],
            "Avløp": [
                f"accum_large-wastewater_tot{par}_tonnes",
                f"accum_spredt_tot{par}_tonnes",
            ],
            "Industri": [f"accum_industry_tot{par}_tonnes"],
            "Bebygd": [f"accum_urban_tot{par}_tonnes"],
            "Bakgrunn": [
                f"accum_agriculture-background_tot{par}_tonnes",
                f"accum_upland_tot{par}_tonnes",
                f"accum_wood_tot{par}_tonnes",
                f"accum_lake_tot{par}_tonnes",
            ],
        }
    else:
        # TOC
        agg_dict = {
            "Akvakultur": [f"accum_aquaculture_to{par}_tonnes"],
            "Jordbruk": [f"accum_agriculture_to{par}_tonnes"],
            "Avløp": [
                f"accum_large-wastewater_to{par}_tonnes",
                f"accum_spredt_to{par}_tonnes",
            ],
            "Industri": [f"accum_industry_to{par}_tonnes"],
            "Bebygd": [f"accum_urban_to{par}_tonnes"],
            "Bakgrunn": [
                f"accum_agriculture-background_to{par}_tonnes",
                f"accum_upland_to{par}_tonnes",
                f"accum_wood_to{par}_tonnes",
            ],
        }

    return agg_dict


def aggregate_parameters(df, par, model):
    """Aggregate columns in TEOTIL output to headings used in the report.

    Args
        df: Dataframe of TEOTIL results
        par: Str. Either 'n' or 'p'
        model: Str. Either 'teotil2' or 'teotil3'

    Returns
        Dataframe.
    """
    agg_dict = get_aggregation_dict_for_columns(par, model=model)
    for group, cols in agg_dict.items():
        df[group] = df[cols].sum(axis=1)

    df = df[["regine", "År"] + list(agg_dict.keys())]

    return df


def aggregate_regions(df, par, out_fold=None):
    """Sum TEOTIL output for the main catchments for each region defined in the
       report.

    Args
        df: Dataframe of results aggregated to the correct column names for the report.
        par: Str. 'n', 'p' or 'c'.
        out_fold: None or str. Default None. Folder to save CSVs to, if desired.

    Returns
        Dict of dataframes. Optionally, results for each region are saved to CSV.
    """
    assert par in ("n", "p", "c")

    result_dict = {}
    for region, catches in REGIONS_DICT.items():
        if len(catches) == 2:
            catch_list = list(range(catches[0], catches[1]))
        else:
            catch_list = list(range(catches[0], catches[1])) + [catches[2]]

        reg_df = df.query("(Vassom in @catch_list) and (Par == @par)").copy()
        reg_df = reg_df.groupby("År").sum(numeric_only=True).reset_index()
        reg_df = reg_df.round(0).astype(int)
        del reg_df["Vassom"]
        result_dict[region] = reg_df

        if out_fold:
            if not os.path.exists(out_fold):
                os.makedirs(out_fold)
            csv_path = os.path.join(out_fold, f"{region}_{par}.csv")
            reg_df.to_csv(csv_path, index=False)

    return result_dict


def make_plot(df, region, par, out_fold):
    """Create a time series plot from a dataframe.

    Args
        df: Dataframe with data for 'region'.
        region: Str. Name of region/heading in report (used as title for plot).
        par: Str. One of ('n', 'p', 'c').
        out_fold: Str. Folder in which to save plot.

    Returns
        None. Plot is saved to 'out_fold'.
    """
    if par == "p":
        ylabel = "Fosfor (tonn)"
    elif par == "n":
        ylabel = "Nitrogen (tonn)"
    elif par == "c":
        ylabel = "Karbon (tonn)"
    else:
        raise ValueError("Could not identify ylabel.")

    df = df.copy()
    ax = df.set_index("År").plot(
        figsize=(8, 6), title=f"{region}: {ylabel.lower()}", xlabel="", fontsize=12
    )
    ax.legend(
        loc="center", bbox_to_anchor=(0.5, -0.2), ncol=4, prop={"size": 12}
    ).get_frame().set_boxstyle("Round", pad=0.2, rounding_size=0.5)
    ax.set_ylabel(ylabel, fontdict={"fontsize": 12, "fontweight": "bold"})
    plt.tight_layout()

    if not os.path.exists(out_fold):
        os.makedirs(out_fold)
    png_path = os.path.join(out_fold, f"{region}_{par}.png")
    plt.savefig(png_path, dpi=200)
    plt.close()


def insert_para_after(para, style="Normal", align="center"):
    """Insert a new paragraph after the given paragraph.

    Args
        para:  Obj. Paragraph object to insert after.
        style: Str. One of Word's pre-defined styles.
        align: Str. One of 'left', 'center' or 'right'.

    Returns
        Paragraph object.
    """
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)

    new_para.style = style

    if align == "center":
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        raise ValueError("'align' must be one of 'left', 'center' or 'right'.")

    return new_para


def insert_image(para, img_path, width_cm=15):
    """Insert an image into a paragraph.

    Args
        para: Obj. Paragraph object to insert after.
        img_path: Str. Path to image.
        width_cm: Float. Width of image in cm.

    Returns
        Run object.
    """
    run = para.add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    return run


def insert_table(doc, para, df, table_style="Grid Table 4 Accent 1"):
    """Insert a dataframe as a table into the given paragraph.

    Args
        doc: Obj. Document object.
        para: Obj. Paragraph object to insert after.
        df: Obj. Dataframe with table data.
        table_style: Str. Word table style to apply. NOTE: If your template document
            doesn't contain any styled tables, the list of available choices may be
            limited. Try

                styles = doc.styles
                table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]

            to see a list of available styles in your document. To add a style, manually
            create a table with the desired style in the template, then save it, delete
            the table and save again. The chosen style should now be available.

    Returns
        None.
    """
    df = df.copy().astype(str)
    df.replace("<NA>", "", inplace=True)

    # Tables can only be added at the end of the document, but they can be moved
    # elsewhere afterwards. See
    # https://github.com/python-openxml/python-docx/issues/156
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = table_style

    # Add header
    for j in range(df.shape[-1]):
        table.cell(0, j).text = df.columns[j]
        table.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add data
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i + 1, j).text = str(df.values[i, j])
            table.cell(i + 1, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Move to desired location (paragraph)
    tbl, p = table._tbl, para._p
    p.addnext(tbl)


def filename_from_heading(heading):
    """Build a file name based on chapter headings in the Word template.

    Args
        heading: Str. Heading from the Word template.

    Returns
        Str. File name for data CSV or plot PNG (without extension).
    """
    name, par = heading.split(":")
    if par[1] == "f":
        par = "p"
    elif par[1] == "n":
        par = "n"
    elif par[1] == "k":
        par = "c"
    else:
        raise ValueError("Could not parse parameter from Word heading.")

    return f"{name}_{par}"


def make_point_map(pt_gdf, title, notes, png_path=None):
    """Make a map of Norway showing point data.

    Args
        pt_gdf: Geodataframe of point data.
        title: Str. Title to use on map.
        notes: Str. Text to add above legend.
        png_path: Str. Default None. Path to save PNG, if desired.

    Returns
        None.
    """
    # Reproject
    pt_gdf = pt_gdf.to_crs(epsg=EPSG)

    # Get countries
    ne_path = shpreader.natural_earth(
        resolution="10m", category="cultural", name="admin_0_countries"
    )
    land_gdf = gpd.read_file(ne_path).query("CONTINENT == 'Europe'").to_crs(epsg=EPSG)

    # Setup plot
    fig, ax = plt.subplots(
        figsize=(8, 11), subplot_kw={"projection": ccrs.UTM(zone=33)}
    )
    ax.set_facecolor("lightblue")

    # Add data
    land_gdf.plot(ax=ax, facecolor="lightyellow", edgecolor="black", linewidth=0.25)
    land_gdf[land_gdf["NAME"] == "Norway"].plot(
        ax=ax, facecolor="palegoldenrod", edgecolor="black", linewidth=0.25
    )
    pt_gdf.plot(
        ax=ax, color="red", edgecolor="black", linewidth=0.25, markersize=15, marker="o"
    )

    # Add a scale bar
    scalebar = ScaleBar(
        1,
        location="lower right",
        scale_loc="top",
        units="m",
        dimension="si-length",
        color="black",
        frameon=False,
        box_color=None,
        font_properties={"size": 10, "weight": "bold", "style": "italic"},
        fixed_value=400,
        fixed_units="km",
        bbox_to_anchor=(0.975, 0.01),
        bbox_transform=ax.transAxes,
    )
    ax.add_artist(scalebar)

    # Set extent
    ax.set_xlim(XMIN, XMAX)
    ax.set_ylim(YMIN, YMAX)

    # Add title
    ax.text(
        0.01,
        0.992,
        title,
        transform=ax.transAxes,
        fontsize=14,
        fontweight="bold",
        verticalalignment="top",
        horizontalalignment="left",
        bbox=dict(facecolor="white", edgecolor="black", boxstyle="square,pad=0.3"),
    )

    # Add notes
    ax.text(
        0.814,
        0.0095,
        notes,
        transform=ax.transAxes,
        fontsize=10,
        fontweight="bold",
        verticalalignment="bottom",
        horizontalalignment="center",
        bbox=dict(facecolor="white", edgecolor="black", boxstyle="square,pad=0.5"),
    )

    # Add axis border
    for spine in ax.spines.values():
        spine.set_edgecolor("black")
        spine.set_linewidth(1)

    # Save
    if png_path:
        plt.savefig(png_path, dpi=200, bbox_inches="tight")


def make_pie_map(gdf, title, offset_dict, png_path=None, pie_scale_factor=3):
    """Make a map of Norway from a polygon geodataframe with pie charts showing
    proportions for each polygon region.

    Args
        gdf: Polygon geodataframe of regional data.
        title: Str. Title to use on map.
        offset_dict: Dict mapping reguon names to tuples describing how much each
            region label should be offset from the centre of each pie chart (in
            map units). E.g. {'Vestland': (-100000, 50000)} would move the 'Vestland'
            label 100 km west and 50 km north.
        png_path: Str. Default None. Path to save PNG, if desired.
        pie_scale_factor: Float. Factor used to scale size of pie charts relative to
            the total for each chart.

    Returns
        None.
    """
    # Sources to include on pies
    src_list = ["Akvakultur", "Jordbruk", "Avløp", "Industri", "Bebygd", "Bakgrunn"]
    colours = [COLOUR_DICT[src] for src in src_list]

    # Reproject
    gdf = gdf.to_crs(epsg=EPSG)

    # Get countries
    ne_path = shpreader.natural_earth(
        resolution="10m", category="cultural", name="admin_0_countries"
    )
    land_gdf = gpd.read_file(ne_path).query("CONTINENT == 'Europe'").to_crs(epsg=EPSG)

    # Setup plot
    fig, ax = plt.subplots(
        figsize=(8, 11), subplot_kw={"projection": ccrs.UTM(zone=33)}
    )
    ax.set_facecolor("lightblue")

    # Add data
    land_gdf.plot(ax=ax, facecolor="lightyellow", edgecolor="black", linewidth=0.25)
    gdf.plot(ax=ax, facecolor="palegoldenrod", edgecolor="black", linewidth=0.25)

    # Set extent
    ax.set_xlim(XMIN, XMAX)
    ax.set_ylim(YMIN, YMAX)

    # Add pie charts
    for idx, row in gdf.iterrows():
        # Get centroid and label details
        reg = row["Område"]
        reg_offset = offset_dict[reg]
        centroid = row["geometry"].centroid
        x, y = centroid.x, centroid.y

        # The pie chart for Glomma needs shifting relative to the centroid.
        # Others seem OK
        if reg == "Glomma":
            x += 30000

        # Get totals for region by source
        values = [row[src] for src in src_list]
        add_pie_chart(fig, ax, values, colours, x, y, pie_scale_factor, reg, reg_offset)

    # Add a scale bar
    scalebar = ScaleBar(
        1,
        location="lower right",
        scale_loc="top",
        units="m",
        dimension="si-length",
        color="black",
        frameon=False,
        box_color=None,
        font_properties={"size": 10, "weight": "bold", "style": "italic"},
        fixed_value=400,
        fixed_units="km",
        bbox_to_anchor=(0.93, 0.01),
        bbox_transform=ax.transAxes,
    )
    ax.add_artist(scalebar)

    # Add title
    ax.text(
        0.01,
        0.992,
        title,
        transform=ax.transAxes,
        fontsize=14,
        fontweight="bold",
        verticalalignment="top",
        horizontalalignment="left",
        bbox=dict(facecolor="white", edgecolor="black", boxstyle="square,pad=0.3"),
    )

    # Add legend
    patches = [
        mpatches.Patch(color=colour, label=label)
        for label, colour in zip(src_list, colours)
    ]
    legend = ax.legend(
        handles=patches,
        loc="lower center",
        bbox_to_anchor=(0.772, 0.1),
        bbox_transform=ax.transAxes,
        frameon=False,
        prop={"size": 10, "weight": "bold"},
    )

    # Get table data
    df = pd.DataFrame(gdf.drop(columns="geometry")).rename(
        columns={"Totalt": "Totalt (tonn)"}
    )
    df = df[["Område", "Totalt (tonn)"]]

    # Use space as 1000s separator
    df["Totalt (tonn)"] = df["Totalt (tonn)"].apply(
        lambda x: "{:,.0f}".format(x).replace(",", " ")
    )

    # Add table
    bbox = Bbox.from_bounds(0.573, 0.3, 0.4, len(df) * 0.03)
    table = ax.table(
        cellText=df.values, colLabels=df.columns, bbox=bbox, cellLoc="right", zorder=10
    )
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1)

    # Make column headings bold and centered
    for key, cell in table.get_celld().items():
        if key[0] == 0:  # Header row
            cell.set_text_props(weight="bold", ha="center", va="center_baseline")
        elif key[1] == 0:  # First column
            cell.set_text_props(ha="center", va="center_baseline")
        elif key[1] == 1:  # Second column
            cell.set_text_props(ha="right", va="center_baseline")

    # Add white border for legend and table area
    rows = 24 if len(df) > 3 else 14
    ax.text(
        0.772,
        0.0095,
        " " * 55 + "\n" * rows + "ETRS89 UTM Sone 33N\n\n\n",
        transform=ax.transAxes,
        fontsize=10,
        fontweight="bold",
        verticalalignment="bottom",
        horizontalalignment="center",
        bbox=dict(facecolor="white", edgecolor="black", boxstyle="square,pad=0.5"),
    )

    # Add axis border
    for spine in ax.spines.values():
        spine.set_edgecolor("black")
        spine.set_linewidth(1)

    # Save
    if png_path:
        plt.savefig(png_path, dpi=200, bbox_inches="tight")


def add_pie_chart(
    fig, ax, values, colors, x, y, pie_scale_factor, region, region_offset
):
    """Add a pie chart to an (x, y) location on a cartographic map.

    Args
        fig: Obj. Figure object to add pie charts to.
        ax: Obj. Axis object on 'fig' to add pie charts to.
        values: List of Float. Values defining the size of each pie slice.
        colours: List of Str. Colour for each element in 'values'. Must be the same
            length as 'values'.
        x: Float. Easting for centre of pie chart in map coordinates.
        y: Float. Northing for centre of pie chart in map coordinates.
        pie_scale_factor: Float. Factor used to scale size of pie charts relative to
            the total for each chart.
        region: Str. Region name to add beside pie chart.
        region_offset: Tuple of Float (x_offset, y_offset). Offset for region label
            relative to pie chart centre in map units.

    Returns
        None.
    """
    bbox = ax.get_position()
    trans = ax.transData.transform((x, y))
    x_fig, y_fig = fig.transFigure.inverted().transform(trans)

    # Scale pie chart size based on total flux
    total = sum(values)
    min_size = 0.03
    width = height = min_size + (total * pie_scale_factor / 1e6)

    # Add axis for pie chart
    pie_ax = fig.add_axes(
        [x_fig - width / 2, y_fig - height / 2, width, height], aspect="equal"
    )

    # Plot pie chart
    pie_ax.pie(
        values,
        colors=colors,
        startangle=90,
        counterclock=False,
        wedgeprops={"edgecolor": "black"},
    )

    # Add region label
    ax.text(
        x + region_offset[0],
        y + region_offset[1],
        region,
        # transform=ax.transAxes,
        fontsize=8,
        fontweight="bold",
        verticalalignment="center",
        horizontalalignment="left",
        bbox=None,
    )